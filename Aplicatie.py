import tkinter as tk
import os
from tkinter import messagebox, TOP, Entry, Label, StringVar
from tkinterdnd2 import *
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches
from datetime import datetime

#-- Global variables --
file_name = " "
entryWidget = None
pathLabel = None

#--  Doc create
if not os.path.isfile("output.docx"):
    document = Document()
    document.save("output.docx")

def add_image_checkbox():
    global entryWidget
    global pathLabel
    global file_name

    if (var1.get()) == 1:
        entryWidget = Entry(window)
        entryWidget.pack(anchor="w", padx=5, pady=5)

        pathLabel = Label(window, text="Drag and drop file in the entry box")
        pathLabel.pack(anchor="w")

        entryWidget.drop_target_register(DND_ALL)
        entryWidget.dnd_bind("<<Drop>>", get_image_path)
        #window.update()

    else:
        entryWidget.pack_forget()
        pathLabel.pack_forget()
        file_name = " "
        #window.update()
'''
def get_file_type(file_path):
    _, file_extension = os.path.splitext(file_path)
    file_type = file_extension.lstrip(".")
    return file_type
'''
def check_file_type(file_path):
    file_extension = os.path.splitext(file_path)[1]
    print(file_extension)
    if (file_extension.lower() == '.apng}'
        or file_extension.lower() == '.avif}'
        or file_extension.lower() == '.gif}'
        or file_extension.lower() == '.png}'
        or file_extension.lower() == '.svg}'
        or file_extension.lower() == '.webp}'
        or file_extension.lower() == '.jpg}'
        or file_extension.lower() == '.jpeg}'
        or file_extension.lower() == '.jfif}'
        or file_extension.lower() == '.pjpeg}'
        or file_extension.lower() == '.pjp}'):
        return True
    else:
        return False

def get_image_path(event):
    #Gets full path
    global file_name
    global file_path
    file_path = event.data
    print (file_path)
    file_name = os.path.basename(file_path)
    file_name = file_name.rstrip("}")

    #check_file_type(file_path)

    if check_file_type(file_path) == True:
        pathLabel.configure(text = event.data)
    else:
        messagebox.showerror("Error", "Wrong file format")

def export_to_word_title(ticket_number, assignee, reviewer, dt_string):

    input_ticket_number= ticket_number
    input_assignee = assignee
    input_reviewer = reviewer
    input_dt_string = dt_string

    document = Document("output.docx")

    total_text_length = sum(len(p.text) for p in document.paragraphs)

    if total_text_length > 0:
        document.add_section(WD_SECTION.NEW_PAGE)

    document.add_paragraph(f"Data: {input_dt_string}")
    document.add_paragraph(f"Assignee: {input_assignee}")
    document.add_paragraph(f"Reviewer: {input_reviewer}")
    document.add_paragraph(f"Ticket number: {input_ticket_number}")

    document.save("output.docx")

def export_to_word():
    global file_name
    input_text1 = text1.get("1.0", "end-1c")
    input_text2 = text2.get("1.0", "end-1c")
    input_text3 = text3.get("1.0", "end-1c")
    '''Adds places holder demo
    def on_entry_click():
        if text1.get("1.0", tk.END).strip() == "Introduceți textul aici":
            text1.delete("1.0", tk.END)
            text1.config(foreground="black")

    def on_focus_out():
        if text1.get("1.0", tk.END).strip() == "":
            text1.insert("1.0", "Introduceți textul aici")
            text1.config(foreground="gray")
    '''
    document = Document("output.docx")
    #document.add_section(WD_SECTION.NEW_PAGE)
    document.add_paragraph(f"Text1: {input_text1}")
    document.add_paragraph(f"Text2: {input_text2}")
    document.add_paragraph(f"Text3: {input_text3}")
    #Add image to word file demo
    document.save("output.docx")
    #add_image_checkbox()
    print(file_name)
    if file_name != " ":
        document = Document("output.docx")
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(file_name)
        document.save("output.docx")

    messagebox.showinfo("Pop-up", "Text was exported to file output.docx")

def open_popup():
    def on_popup_close():
        ticket_number = entry1.get()
        assignee = entry2.get()
        reviewer = entry3.get()

        now = datetime.now()
        dt_string = now.strftime("%d %B %Y %H:%M:%S")

        export_to_word_title(ticket_number, assignee, reviewer, dt_string)
        popup.destroy()

    popup = tk.Toplevel()
    popup.geometry("400x200")
    popup.title("Pop-up")

    popup_label = tk.Label(popup, text="Number ticket:")
    popup_label.pack()

    entry1 = tk.Entry(popup)
    entry1.pack()
    entry1.focus()
    entry1.insert(0, "GMCTC-")

    popup_label = tk.Label(popup, text="Assignee:")
    popup_label.pack()

    entry2 = tk.Entry(popup)
    entry2.pack()

    popup_label = tk.Label(popup, text="Reviewer:")
    popup_label.pack()

    entry3 = tk.Entry(popup)
    entry3.pack()

    popup_button = tk.Button(popup, text="Export to Word", command=on_popup_close)
    popup_button.pack()

#-- main --
#window = tk.Tk()
window = TkinterDnD.TixTk()
window.geometry("500x500")
#window.resizable(False, False)
var1 = tk.IntVar()

label = tk.Label(window, text="Remark found in variant: ")
label.pack(anchor="w")

text1 = tk.Text(window, width=0,height=0,borderwidth =7)
text1.pack(anchor="w",fill="x")

label = tk.Label(window, text="Remark found in file:")
label.pack(anchor="w")

text2 = tk.Text(window, width=0,height=0,borderwidth =7)
text2.pack(anchor="w",fill="x")

label = tk.Label(window, text="Remark")
label.pack(anchor="w")

text3 = tk.Text(window, width=0,height=0,borderwidth =7)
text3.pack(anchor="w",fill="both", expand=True)

checkbox = tk.Checkbutton(window, text='Add Image',variable=var1, onvalue=1, offvalue=0, command = add_image_checkbox)
checkbox.pack(anchor ="w", padx= 0, pady=0)

button = tk.Button(window, text="Export to Word", command=export_to_word)
button.pack(anchor="c")

popup_button = tk.Button(window, text="Open Pop-up", command=open_popup)
popup_button.pack(anchor="c")

window.mainloop()